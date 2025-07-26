# SPDX-License-Identifier: MIT
# Copyright (c) 2025 dr.max

import atexit
import os
import signal
import tempfile
import warnings
from typing import Any

import docx
import PyPDF2
from fastapi import FastAPI, File, UploadFile
from pydantic import BaseModel

# Suppress Pydantic deprecation and schema warnings from dependencies
warnings.filterwarnings(
    "ignore", category=DeprecationWarning, message=".*PydanticDeprecatedSince211.*"
)
warnings.filterwarnings(
    "ignore", category=UserWarning, message=".*PydanticJsonSchemaWarning.*"
)
warnings.filterwarnings("ignore", message=".*model_fields.*")
warnings.filterwarnings("ignore", message=".*not JSON serializable.*")

# Suppress ResourceWarnings from dependencies
warnings.filterwarnings("ignore", category=ResourceWarning, message=".*unclosed.*")
warnings.filterwarnings(
    "ignore", category=ResourceWarning, message=".*Enable tracemalloc.*"
)
warnings.filterwarnings(
    "ignore", category=ResourceWarning
)  # General ResourceWarning suppression

app = FastAPI(title="RagMe MCP Server")


# Cleanup function
def cleanup():
    """Clean up resources when the application shuts down."""
    try:
        # Clean up any temporary files that might still exist
        temp_dir = tempfile.gettempdir()
        for filename in os.listdir(temp_dir):
            if filename.startswith("tmp") and (
                filename.endswith(".pdf") or filename.endswith(".docx")
            ):
                try:
                    os.unlink(os.path.join(temp_dir, filename))
                except Exception:
                    pass
    except Exception as e:
        print(f"Error during MCP cleanup: {e}")


# Register cleanup handlers
atexit.register(cleanup)


def signal_handler(signum, frame):
    """Handle shutdown signals."""
    cleanup()
    # Don't call sys.exit(0) as it causes asyncio errors


# Register signal handlers
signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)


class ToolResponse(BaseModel):
    success: bool
    data: dict[str, Any] | None = None
    error: str | None = None


class Base64FileRequest(BaseModel):
    filename: str
    content: str  # base64 encoded content
    content_type: str


@app.post("/tool/process_pdf", response_model=ToolResponse)
async def process_pdf(file: UploadFile = File(...)):
    """Process a PDF file uploaded via multipart form data"""
    try:
        # Validate file type
        if not file.filename.lower().endswith(".pdf"):
            return ToolResponse(success=False, error="File must be a PDF")

        # Read the uploaded file content
        content = await file.read()

        # Create a temporary file to store the PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the PDF using PyPDF2
            with open(temp_file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                # Extract text from all pages
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                # Get metadata and ensure filename is preserved
                pdf_metadata = pdf_reader.metadata or {}
                metadata = {
                    "filename": file.filename,  # Always preserve the filename
                    **pdf_metadata,  # Include PDF metadata but don't let it override filename
                }

                return ToolResponse(
                    success=True,
                    data={
                        "data": {
                            "filename": file.filename,
                            "text": text,
                            "page_count": len(pdf_reader.pages),
                        },
                        "metadata": metadata,
                    },
                )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_docx", response_model=ToolResponse)
async def process_docx(file: UploadFile = File(...)):
    """Process a DOCX file uploaded via multipart form data"""
    try:
        # Validate file type
        if not file.filename.lower().endswith(".docx"):
            return ToolResponse(success=False, error="File must be a DOCX")

        # Read the uploaded file content
        content = await file.read()

        # Create a temporary file to store the DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the DOCX using python-docx
            doc = docx.Document(temp_file_path)

            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Get core properties and ensure filename is preserved
            core_props = doc.core_properties
            metadata = {
                "filename": file.filename,  # Always preserve the filename
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": file.filename,
                        "text": text,
                        "tables": tables,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    },
                    "metadata": metadata,
                },
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_pdf_base64", response_model=ToolResponse)
async def process_pdf_base64(request: Base64FileRequest):
    try:
        # Validate file type
        if not request.filename.lower().endswith(".pdf"):
            return ToolResponse(success=False, error="File must be a PDF")

        # Decode base64 content
        import base64

        content = base64.b64decode(request.content)

        # Create a temporary file to store the decoded PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the PDF using PyPDF2
            with open(temp_file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                # Extract text from all pages
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                # Get metadata and ensure filename is preserved
                pdf_metadata = pdf_reader.metadata or {}
                metadata = {
                    "filename": request.filename,  # Always preserve the filename
                    **pdf_metadata,  # Include PDF metadata but don't let it override filename
                }

                return ToolResponse(
                    success=True,
                    data={
                        "data": {
                            "filename": request.filename,
                            "text": text,
                            "page_count": len(pdf_reader.pages),
                        },
                        "metadata": metadata,
                    },
                )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_docx_base64", response_model=ToolResponse)
async def process_docx_base64(request: Base64FileRequest):
    try:
        # Validate file type
        if not request.filename.lower().endswith(".docx"):
            return ToolResponse(success=False, error="File must be a DOCX")

        # Decode base64 content
        import base64

        content = base64.b64decode(request.content)

        # Create a temporary file to store the decoded DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the DOCX using python-docx
            doc = docx.Document(temp_file_path)

            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Get core properties and ensure filename is preserved
            core_props = doc.core_properties
            metadata = {
                "filename": request.filename,  # Always preserve the filename
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": request.filename,
                        "text": text,
                        "tables": tables,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    },
                    "metadata": metadata,
                },
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8022)
