# SPDX-License-Identifier: MIT
# Copyright (c) 2025 dr.max

import atexit
import logging
import os
import signal
import tempfile
import warnings
from typing import Any

import docx
import PyPDF2
from fastapi import FastAPI, File, UploadFile
from pydantic import BaseModel

from ..utils.config_manager import config
from ..utils.storage import StorageService

# Try to import additional PDF processing libraries
try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available, falling back to PyPDF2")

try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available, falling back to PyPDF2")

# Suppress Pydantic deprecation and schema warnings from dependencies
warnings.filterwarnings(
    "ignore", category=DeprecationWarning, message=".*PydanticDeprecatedSince211.*"
)
warnings.filterwarnings(
    "ignore", category=UserWarning, message=".*PydanticJsonSchemaWarning.*"
)
warnings.filterwarnings("ignore", message=".*model_fields.*")
warnings.filterwarnings("ignore", message=".*not JSON serializable.*")
warnings.filterwarnings(
    "ignore", category=DeprecationWarning, message=".*PydanticDeprecatedSince20.*"
)
warnings.filterwarnings(
    "ignore", category=DeprecationWarning, message=".*class-based `config`.*"
)

# Get application configuration
app_config = config.get("application", {})
app = FastAPI(
    title=f"{app_config.get('name', 'RAGme')} MCP Server",
    description="MCP (Model Context Protocol) Server for document processing",
    version=app_config.get("version", "1.0.0"),
)

# Storage service will be initialized lazily when needed
_storage_service = None


def get_storage_service():
    """Get or create storage service instance"""
    global _storage_service
    if _storage_service is None:
        _storage_service = StorageService(config)
    return _storage_service


# Cleanup function
def cleanup():
    """Clean up resources when the application shuts down."""
    try:
        # Clean up any temporary files that might still exist
        temp_dir = tempfile.gettempdir()
        for filename in os.listdir(temp_dir):
            if filename.startswith("tmp") and (
                filename.endswith(".pdf") or filename.endswith(".docx")
            ):
                try:
                    os.unlink(os.path.join(temp_dir, filename))
                except Exception:
                    pass
    except Exception as e:
        print(f"Error during MCP cleanup: {e}")


# Register cleanup handlers
atexit.register(cleanup)


def signal_handler(signum, frame):
    """Handle shutdown signals."""
    cleanup()
    # Don't call sys.exit(0) as it causes asyncio errors


# Register signal handlers
signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)


class ToolResponse(BaseModel):
    success: bool
    data: dict[str, Any] | None = None
    error: str | None = None


class Base64FileRequest(BaseModel):
    filename: str
    content: str  # base64 encoded content
    content_type: str


def process_pdf_with_fallback(pdf_path: str) -> tuple[str, int, dict]:
    """
    Process PDF using multiple libraries as fallbacks to handle corrupted PDFs.
    Returns (text, page_count, metadata)
    """
    errors = []

    # Try PyMuPDF first (most robust for corrupted PDFs)
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            metadata = doc.metadata
            page_count = len(doc)
            doc.close()
            logging.info(f"Successfully processed PDF with PyMuPDF: {pdf_path}")
            return text, page_count, metadata
        except Exception as e:
            error_msg = f"PyMuPDF failed: {str(e)}"
            errors.append(error_msg)
            logging.warning(error_msg)

    # Try pdfplumber (good for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                metadata = pdf.metadata
                page_count = len(pdf.pages)
            logging.info(f"Successfully processed PDF with pdfplumber: {pdf_path}")
            return text, page_count, metadata
        except Exception as e:
            error_msg = f"pdfplumber failed: {str(e)}"
            errors.append(error_msg)
            logging.warning(error_msg)

    # Try PyPDF2 (original method)
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            metadata = pdf_reader.metadata or {}
            page_count = len(pdf_reader.pages)
        logging.info(f"Successfully processed PDF with PyPDF2: {pdf_path}")
        return text, page_count, metadata
    except Exception as e:
        error_msg = f"PyPDF2 failed: {str(e)}"
        errors.append(error_msg)
        logging.error(error_msg)

    # If all methods fail, try to repair the PDF
    try:
        # Try to read the file as binary and look for PDF structure
        with open(pdf_path, "rb") as f:
            content = f.read()

        # Look for PDF header
        if content.startswith(b"%PDF"):
            # Try to find the end of the PDF
            end_marker = b"%%EOF"
            if end_marker in content:
                # Try to truncate at the last %%EOF
                last_eof = content.rfind(end_marker)
                if last_eof > 0:
                    # Create a temporary repaired file
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".pdf"
                    ) as temp_file:
                        temp_file.write(content[: last_eof + len(end_marker)])
                        temp_path = temp_file.name

                    try:
                        # Try PyMuPDF again with repaired file
                        if PYMUPDF_AVAILABLE:
                            doc = fitz.open(temp_path)
                            text = ""
                            for page in doc:
                                text += page.get_text() + "\n"
                            metadata = doc.metadata
                            page_count = len(doc)
                            doc.close()
                            logging.info(
                                f"Successfully processed repaired PDF with PyMuPDF: {pdf_path}"
                            )
                            return text, page_count, metadata
                    finally:
                        os.unlink(temp_path)

        # If repair attempt fails, return partial text
        logging.warning(f"All PDF processing methods failed for {pdf_path}")
        return f"[PDF processing failed: {'; '.join(errors)}]", 0, {}

    except Exception as e:
        logging.error(f"PDF repair attempt failed: {str(e)}")
        return f"[PDF processing failed: {'; '.join(errors)}]", 0, {}


@app.post("/tool/process_pdf", response_model=ToolResponse)
async def process_pdf(file: UploadFile = File(...)):
    """Process a PDF file uploaded via multipart form data"""
    try:
        # Validate file type
        if not file.filename.lower().endswith(".pdf"):
            return ToolResponse(success=False, error="File must be a PDF")

        # Read the uploaded file content
        content = await file.read()

        # Create a temporary file to store the PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the PDF using robust fallback method
            text, page_count, pdf_metadata = process_pdf_with_fallback(temp_file_path)

            # Get metadata and ensure filename is preserved
            metadata = {
                "filename": file.filename,  # Always preserve the filename
                **(
                    pdf_metadata or {}
                ),  # Include PDF metadata but don't let it override filename
            }

            # Copy file to storage if enabled
            storage_path = None
            if config.is_copy_uploaded_docs_enabled():
                try:
                    from datetime import datetime

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                    storage_path = f"documents/{timestamp}_{file.filename}"
                    get_storage_service().upload_data(
                        data=content,
                        object_name=storage_path,
                        content_type="application/pdf",
                    )
                    print(f"Copied PDF to storage: {storage_path}")
                except Exception as storage_error:
                    print(f"Failed to copy PDF to storage: {storage_error}")
                    storage_path = None

            # Add storage path to metadata if file was copied to storage
            if storage_path:
                metadata["storage_path"] = storage_path

            # Extract and process images from PDF if enabled
            extracted_images_count = 0
            if config.get("pdf_image_extraction", {}).get("enabled", True):
                try:
                    from ..utils.pdf_image_extractor import pdf_image_extractor

                    # Extract images from the PDF
                    extracted_images = pdf_image_extractor.extract_images_from_pdf(
                        temp_file_path, file.filename, storage_path
                    )

                    if extracted_images:
                        # Add extracted images to the image collection
                        success = (
                            pdf_image_extractor.add_extracted_images_to_collection(
                                extracted_images
                            )
                        )
                        if success:
                            extracted_images_count = len(extracted_images)
                            print(
                                f"Successfully extracted and processed {extracted_images_count} images from PDF"
                            )
                        else:
                            print("Failed to add extracted images to collection")
                    else:
                        print("No images found in PDF")

                except Exception as image_error:
                    print(f"Error extracting images from PDF: {image_error}")
                    # Don't fail the entire PDF processing if image extraction fails

            # Add extracted images count to metadata
            metadata["extracted_images_count"] = extracted_images_count

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": file.filename,
                        "text": text,
                        "page_count": page_count,
                        "extracted_images_count": extracted_images_count,
                    },
                    "metadata": metadata,
                },
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_docx", response_model=ToolResponse)
async def process_docx(file: UploadFile = File(...)):
    """Process a DOCX file uploaded via multipart form data"""
    try:
        # Validate file type
        if not file.filename.lower().endswith(".docx"):
            return ToolResponse(success=False, error="File must be a DOCX")

        # Read the uploaded file content
        content = await file.read()

        # Create a temporary file to store the DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the DOCX using python-docx
            doc = docx.Document(temp_file_path)

            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Get core properties and ensure filename is preserved
            core_props = doc.core_properties
            metadata = {
                "filename": file.filename,  # Always preserve the filename
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            # Copy file to storage if enabled
            storage_path = None
            if config.is_copy_uploaded_docs_enabled():
                try:
                    from datetime import datetime

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                    storage_path = f"documents/{timestamp}_{file.filename}"
                    get_storage_service().upload_data(
                        data=content,
                        object_name=storage_path,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    print(f"Copied DOCX to storage: {storage_path}")
                except Exception as storage_error:
                    print(f"Failed to copy DOCX to storage: {storage_error}")
                    storage_path = None

            # Add storage path to metadata if file was copied to storage
            if storage_path:
                metadata["storage_path"] = storage_path

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": file.filename,
                        "text": text,
                        "tables": tables,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    },
                    "metadata": metadata,
                },
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_pdf_base64", response_model=ToolResponse)
async def process_pdf_base64(request: Base64FileRequest):
    try:
        # Validate file type
        if not request.filename.lower().endswith(".pdf"):
            return ToolResponse(success=False, error="File must be a PDF")

        # Decode base64 content
        import base64

        content = base64.b64decode(request.content)

        # Create a temporary file to store the decoded PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the PDF using PyPDF2
            with open(temp_file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                # Extract text from all pages
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                # Get metadata and ensure filename is preserved
                pdf_metadata = pdf_reader.metadata or {}
                metadata = {
                    "filename": request.filename,  # Always preserve the filename
                    **pdf_metadata,  # Include PDF metadata but don't let it override filename
                }

                # Copy file to storage if enabled
                storage_path = None
                if config.is_copy_uploaded_docs_enabled():
                    try:
                        from datetime import datetime

                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                        storage_path = f"documents/{timestamp}_{request.filename}"
                        get_storage_service().upload_data(
                            data=content,
                            object_name=storage_path,
                            content_type="application/pdf",
                        )
                        print(f"Copied PDF to storage: {storage_path}")
                    except Exception as storage_error:
                        print(f"Failed to copy PDF to storage: {storage_error}")
                        storage_path = None

                # Add storage path to metadata if file was copied to storage
                if storage_path:
                    metadata["storage_path"] = storage_path

                # Extract and process images from PDF if enabled
                extracted_images_count = 0
                if config.get("pdf_image_extraction", {}).get("enabled", True):
                    try:
                        from ..utils.pdf_image_extractor import pdf_image_extractor

                        # Extract images from the PDF
                        extracted_images = pdf_image_extractor.extract_images_from_pdf(
                            temp_file_path, request.filename, storage_path
                        )

                        if extracted_images:
                            # Add extracted images to the image collection
                            success = (
                                pdf_image_extractor.add_extracted_images_to_collection(
                                    extracted_images
                                )
                            )
                            if success:
                                extracted_images_count = len(extracted_images)
                                print(
                                    f"Successfully extracted and processed {extracted_images_count} images from PDF"
                                )
                            else:
                                print("Failed to add extracted images to collection")
                        else:
                            print("No images found in PDF")

                    except Exception as image_error:
                        print(f"Error extracting images from PDF: {image_error}")
                        # Don't fail the entire PDF processing if image extraction fails

                return ToolResponse(
                    success=True,
                    data={
                        "data": {
                            "filename": request.filename,
                            "text": text,
                            "page_count": len(pdf_reader.pages),
                            "extracted_images_count": extracted_images_count,
                        },
                        "metadata": metadata,
                    },
                )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


@app.post("/tool/process_docx_base64", response_model=ToolResponse)
async def process_docx_base64(request: Base64FileRequest):
    try:
        # Validate file type
        if not request.filename.lower().endswith(".docx"):
            return ToolResponse(success=False, error="File must be a DOCX")

        # Decode base64 content
        import base64

        content = base64.b64decode(request.content)

        # Create a temporary file to store the decoded DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the DOCX using python-docx
            doc = docx.Document(temp_file_path)

            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Get core properties and ensure filename is preserved
            core_props = doc.core_properties
            metadata = {
                "filename": request.filename,  # Always preserve the filename
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            # Copy file to storage if enabled
            storage_path = None
            if config.is_copy_uploaded_docs_enabled():
                try:
                    from datetime import datetime

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                    storage_path = f"documents/{timestamp}_{request.filename}"
                    get_storage_service().upload_data(
                        data=content,
                        object_name=storage_path,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    print(f"Copied DOCX to storage: {storage_path}")
                except Exception as storage_error:
                    print(f"Failed to copy DOCX to storage: {storage_error}")
                    storage_path = None

            # Add storage path to metadata if file was copied to storage
            if storage_path:
                metadata["storage_path"] = storage_path

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": request.filename,
                        "text": text,
                        "tables": tables,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    },
                    "metadata": metadata,
                },
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(success=False, error=str(e))


if __name__ == "__main__":
    import uvicorn

    # Get network configuration
    network_config = config.get_network_config()
    mcp_config = network_config.get("mcp", {})
    host = mcp_config.get("host", "0.0.0.0")
    port = mcp_config.get("port", 8022)

    uvicorn.run(app, host=host, port=port)
