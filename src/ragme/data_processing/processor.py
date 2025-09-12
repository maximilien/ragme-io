# SPDX-License-Identifier: MIT
# Copyright (c) 2025 dr.max

import json
import os
import tempfile
import time
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
import fitz  # PyMuPDF
import PyPDF2
from docx import Document

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

from ..utils.config_manager import config
from ..utils.image_processor import image_processor
from ..vdbs.vector_db_factory import create_vector_database


class DocumentProcessor:
    """
    Core document processor for the RAGme data processing pipeline.
    Handles individual document processing including text extraction,
    chunking, image processing, and VDB storage.
    """

    def __init__(self, batch_size: int = 3, retry_limit: int = 3):
        """
        Initialize the document processor.

        Args:
            batch_size: Number of parallel processes (for future use)
            retry_limit: Maximum number of retry attempts per document
        """
        self.batch_size = batch_size
        self.retry_limit = retry_limit

        # Get VDB configuration
        self.text_collection_name = config.get_text_collection_name()
        self.image_collection_name = config.get_image_collection_name()

        # Create VDB instances
        self.text_vdb = create_vector_database(collection_name=self.text_collection_name)
        self.image_vdb = create_vector_database(collection_name=self.image_collection_name)

        # Ensure collections are set up
        self.text_vdb.setup()
        self.image_vdb.setup()

        # Get chunking configuration from VDB settings
        db_config = config.get_database_config()
        if db_config:
            self.chunk_size = db_config.get("chunk_size", 1000)
            self.chunk_overlap = db_config.get("chunk_overlap", 100)
            self.chunk_overlap_ratio = db_config.get("chunk_overlap_ratio", 0.2)
        else:
            self.chunk_size = 1000
            self.chunk_overlap = 100
            self.chunk_overlap_ratio = 0.2

        # Supported file extensions
        self.supported_document_extensions = {".pdf", ".docx"}
        self.supported_image_extensions = {
            ".jpg",
            ".jpeg",
            ".png",
            ".gif",
            ".webp",
            ".bmp",
            ".heic",
            ".heif",
            ".tiff",
            ".tif"
        }

    def is_supported_file(self, file_path: str) -> bool:
        """Check if a file is supported for processing."""
        file_ext = Path(file_path).suffix.lower()
        return (
            file_ext in self.supported_document_extensions
            or file_ext in self.supported_image_extensions
        )

    def get_file_type(self, file_path: str) -> Optional[str]:
        """Get the type of file (document or image)."""
        file_ext = Path(file_path).suffix.lower()
        if file_ext in self.supported_document_extensions:
            return "document"
        elif file_ext in self.supported_image_extensions:
            return "image"
        return None

    def process_pdf_with_fallback(self, pdf_path: str) -> Tuple[str, int, dict, List[str]]:
        """
        Process PDF using multiple libraries as fallbacks to handle corrupted PDFs.
        Returns (text, page_count, metadata, extracted_image_paths)
        """
        errors = []
        extracted_images = []

        # Try PyMuPDF first (most robust for corrupted PDFs)
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            # Extract text from all pages
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text() + "\n"
                
                # Extract images from this page if PDF image extraction is enabled
                if config.get("pdf_image_extraction.enabled", True):
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            # Skip images that are too small or too large
                            min_size_kb = config.get("pdf_image_extraction.min_image_size_kb", 1)
                            max_size_mb = config.get("pdf_image_extraction.max_image_size_mb", 10)
                            
                            if pix.n < 5:  # Skip if not a good image format
                                img_data = pix.tobytes("png")
                                img_size_kb = len(img_data) / 1024
                                
                                if min_size_kb <= img_size_kb <= (max_size_mb * 1024):
                                    # Save extracted image to temp file
                                    temp_img_path = tempfile.mktemp(suffix=f"_p{page_num}_i{img_index}.png")
                                    with open(temp_img_path, "wb") as img_file:
                                        img_file.write(img_data)
                                    extracted_images.append(temp_img_path)
                                    
                            pix = None  # Clean up
                        except Exception as e:
                            print(f"Warning: Failed to extract image {img_index} from page {page_num}: {e}")
            
            metadata = doc.metadata
            page_count = len(doc)
            doc.close()
            return text, page_count, metadata, extracted_images
            
        except Exception as e:
            error_msg = f"PyMuPDF failed: {str(e)}"
            errors.append(error_msg)

        # Try pdfplumber (good for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    metadata = pdf.metadata
                    page_count = len(pdf.pages)
                return text, page_count, metadata, extracted_images
                
            except Exception as e:
                error_msg = f"pdfplumber failed: {str(e)}"
                errors.append(error_msg)

        # Try PyPDF2 (fallback)
        try:
            with open(pdf_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                metadata = pdf_reader.metadata or {}
                page_count = len(pdf_reader.pages)
            return text, page_count, metadata, extracted_images
            
        except Exception as e:
            error_msg = f"PyPDF2 failed: {str(e)}"
            errors.append(error_msg)

        # If all methods failed, return error
        return f"[PDF processing failed: {'; '.join(errors)}]", 0, {}, extracted_images

    def process_docx_file(self, docx_path: str) -> Tuple[str, dict]:
        """
        Process DOCX file and extract text and metadata.
        Returns (text, metadata)
        """
        try:
            doc = Document(docx_path)
            
            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            
            # Get core properties
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "tables_count": len(tables),
                "paragraph_count": len(doc.paragraphs),
            }
            
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Failed to process DOCX file: {str(e)}")

    def chunk_text(self, text: str, filename: str = "unknown") -> List[str]:
        """
        Split text into chunks with overlap for better context retention.
        
        Args:
            text: Text to chunk
            filename: Name of the source file for metadata
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0
        overlap_size = int(self.chunk_size * self.chunk_overlap_ratio)

        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings followed by whitespace
                for i in range(end, max(start + self.chunk_size // 2, start), -1):
                    if text[i] in ".!?" and i + 1 < len(text) and text[i + 1].isspace():
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            if end >= len(text):
                break
            start = end - overlap_size

        return chunks

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a single document file (PDF or DOCX).
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Processing results dictionary
        """
        start_time = time.time()
        file_name = Path(file_path).name
        file_size_kb = os.path.getsize(file_path) / 1024
        file_ext = Path(file_path).suffix.lower()
        
        results = {
            "file_name": file_name,
            "file_path": file_path,
            "file_size_kb": file_size_kb,
            "file_type": "document",
            "document_type": file_ext[1:],  # Remove the dot
            "processing_start_time": datetime.fromtimestamp(start_time).isoformat(),
            "chunks": [],
            "extracted_images": [],
            "errors": [],
            "metadata": {},
            "timing": {},
        }
        
        try:
            # Process based on file type
            if file_ext == ".pdf":
                extract_start = time.time()
                text, page_count, pdf_metadata, extracted_image_paths = self.process_pdf_with_fallback(file_path)
                extract_time = time.time() - extract_start
                
                results["metadata"].update({
                    "page_count": page_count,
                    **pdf_metadata,
                })
                results["timing"]["text_extraction"] = extract_time
                
                # Process extracted images
                image_processing_start = time.time()
                for img_path in extracted_image_paths:
                    try:
                        img_result = self.process_image(img_path, is_extracted=True, source_document=file_name)
                        results["extracted_images"].append(img_result)
                    except Exception as e:
                        results["errors"].append(f"Failed to process extracted image {img_path}: {str(e)}")
                    finally:
                        # Clean up temporary image file
                        if os.path.exists(img_path):
                            os.unlink(img_path)
                
                results["timing"]["image_processing"] = time.time() - image_processing_start
                
            elif file_ext == ".docx":
                extract_start = time.time()
                text, docx_metadata = self.process_docx_file(file_path)
                extract_time = time.time() - extract_start
                
                results["metadata"].update(docx_metadata)
                results["timing"]["text_extraction"] = extract_time
            
            else:
                raise ValueError(f"Unsupported document type: {file_ext}")
            
            # Chunk the text
            chunk_start = time.time()
            chunks = self.chunk_text(text, file_name)
            chunk_time = time.time() - chunk_start
            
            results["timing"]["chunking"] = chunk_time
            results["chunk_count"] = len(chunks)
            results["average_chunk_size_kb"] = sum(len(chunk) for chunk in chunks) / len(chunks) / 1024 if chunks else 0
            
            # Store chunks in VDB
            store_start = time.time()
            self._store_document_chunks(file_path, file_name, chunks, results["metadata"])
            store_time = time.time() - store_start
            
            results["timing"]["vdb_storage"] = store_time
            results["chunks"] = [{"index": i, "size_bytes": len(chunk)} for i, chunk in enumerate(chunks)]
            
        except Exception as e:
            error_msg = f"Failed to process document: {str(e)}"
            results["errors"].append(error_msg)
            print(f"Error processing {file_name}: {error_msg}")
            print(traceback.format_exc())
        
        # Calculate total processing time
        total_time = time.time() - start_time
        results["timing"]["total"] = total_time
        results["processing_end_time"] = datetime.now().isoformat()
        
        return results

    def process_image(self, file_path: str, is_extracted: bool = False, source_document: str = None) -> Dict[str, Any]:
        """
        Process a single image file.
        
        Args:
            file_path: Path to the image file
            is_extracted: Whether this image was extracted from a document
            source_document: Name of source document if extracted
            
        Returns:
            Processing results dictionary
        """
        start_time = time.time()
        file_name = Path(file_path).name
        file_size_kb = os.path.getsize(file_path) / 1024
        
        results = {
            "file_name": file_name,
            "file_path": file_path,
            "file_size_kb": file_size_kb,
            "file_type": "image",
            "is_extracted": is_extracted,
            "source_document": source_document,
            "processing_start_time": datetime.fromtimestamp(start_time).isoformat(),
            "errors": [],
            "metadata": {},
            "timing": {},
        }
        
        try:
            # Process image with full pipeline
            process_start = time.time()
            image_url = f"file://{file_path}"
            processed_image = image_processor.process_image(image_url)
            process_time = time.time() - process_start
            
            results["timing"]["image_processing"] = process_time
            
            # Extract individual timing for different phases
            exif_success = "exif" in processed_image and processed_image.get("exif")
            classification_success = "classification" in processed_image and not processed_image["classification"].get("error")
            ocr_success = "ocr_content" in processed_image and processed_image["ocr_content"].get("ocr_processing", False)
            
            results["exif_extracted"] = exif_success
            results["ai_classification_features"] = len(processed_image.get("classification", {}).get("classifications", []))
            results["ocr_success"] = ocr_success
            results["ocr_text_length"] = len(processed_image.get("ocr_content", {}).get("extracted_text", ""))
            
            # Prepare metadata for VDB storage
            combined_metadata = {
                "filename": file_name,
                "file_size": file_size_kb,
                "content_type": "image",
                "date_added": datetime.now().isoformat(),
                "processed_by": "RAGme document processing pipeline",
                "processing_time": process_time,
                "is_extracted_from_document": is_extracted,
                "source_document": source_document,
                **processed_image.get("exif", {}),
                "classification": processed_image.get("classification", {}),
                "ocr_content": processed_image.get("ocr_content", {}),
            }
            
            results["metadata"] = combined_metadata
            
            # Store in VDB
            store_start = time.time()
            self._store_image_in_vdb(image_url, file_name, processed_image, combined_metadata)
            store_time = time.time() - store_start
            
            results["timing"]["vdb_storage"] = store_time
            
        except Exception as e:
            error_msg = f"Failed to process image: {str(e)}"
            results["errors"].append(error_msg)
            print(f"Error processing image {file_name}: {error_msg}")
            print(traceback.format_exc())
        
        # Calculate total processing time
        total_time = time.time() - start_time
        results["timing"]["total"] = total_time
        results["processing_end_time"] = datetime.now().isoformat()
        
        return results

    def _store_document_chunks(self, file_path: str, filename: str, chunks: List[str], metadata: Dict[str, Any]):
        """Store document chunks in the text VDB collection."""
        unique_url = f"file://{file_path}"
        
        if len(chunks) == 1:
            # Single chunk - store as regular document
            combined_metadata = {
                **metadata,
                "filename": filename,
                "content_type": "document",
                "date_added": datetime.now().isoformat(),
                "processed_by": "RAGme document processing pipeline",
                "processing_time": metadata.get("processing_time", 0),
            }
            
            documents = [{
                "text": chunks[0],
                "url": unique_url,
                "metadata": combined_metadata,
            }]
        else:
            # Multiple chunks - store each chunk as separate document
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **metadata,
                    "filename": filename,
                    "content_type": "document",
                    "date_added": datetime.now().isoformat(),
                    "processed_by": "RAGme document processing pipeline",
                    "processing_time": metadata.get("processing_time", 0),
                    "total_chunks": len(chunks),
                    "is_chunked": True,
                    "chunk_index": i,
                    "chunk_sizes": [len(chunk) for chunk in chunks],
                    "original_filename": filename,
                }
                
                documents.append({
                    "text": chunk,
                    "url": f"{unique_url}#chunk-{i}",
                    "metadata": chunk_metadata,
                })
        
        # Write to VDB
        self.text_vdb.write_documents(documents)

    def _store_image_in_vdb(self, image_url: str, filename: str, processed_image: Dict[str, Any], metadata: Dict[str, Any]):
        """Store image in the image VDB collection."""
        # Check if VDB supports images directly
        if self.image_vdb.supports_images():
            # Store as image with metadata
            self.image_vdb.write_images([{
                "url": image_url,
                "metadata": metadata,
            }])
        else:
            # Fallback: store as text document with image metadata
            top_pred = processed_image.get("classification", {}).get("top_prediction", {})
            label = top_pred.get("label", "unknown")
            
            # Include OCR content if available
            ocr_content = processed_image.get("ocr_content", {})
            ocr_text = ocr_content.get("extracted_text", "") if ocr_content else ""
            
            text_representation = f"Image: {filename}\nClassification: {label}\n"
            
            if ocr_text:
                text_representation += f"OCR Content: {ocr_text}\n"
            
            text_representation += f"Metadata: {str(metadata)}"
            
            self.image_vdb.write_documents([{
                "url": image_url,
                "text": text_representation,
                "metadata": metadata,
            }])

    def process_file_with_retry(self, file_path: str, max_retries: int = None) -> Dict[str, Any]:
        """
        Process a file with retry logic.
        
        Args:
            file_path: Path to the file to process
            max_retries: Maximum number of retry attempts (uses instance default if None)
            
        Returns:
            Processing results dictionary
        """
        if max_retries is None:
            max_retries = self.retry_limit
        
        file_type = self.get_file_type(file_path)
        if not file_type:
            return {
                "file_name": Path(file_path).name,
                "file_path": file_path,
                "file_type": "unsupported",
                "errors": [f"Unsupported file type: {Path(file_path).suffix}"],
                "success": False,
                "retry_count": 0,
            }
        
        last_error = None
        for attempt in range(max_retries + 1):
            try:
                if file_type == "document":
                    result = self.process_document(file_path)
                else:  # file_type == "image"
                    result = self.process_image(file_path)
                
                result["success"] = len(result.get("errors", [])) == 0
                result["retry_count"] = attempt
                return result
                
            except Exception as e:
                last_error = str(e)
                if attempt < max_retries:
                    print(f"Attempt {attempt + 1} failed for {Path(file_path).name}: {last_error}. Retrying...")
                    time.sleep(1)  # Brief pause before retry
                else:
                    print(f"All {max_retries + 1} attempts failed for {Path(file_path).name}: {last_error}")
        
        # All retries failed
        return {
            "file_name": Path(file_path).name,
            "file_path": file_path,
            "file_type": file_type,
            "errors": [f"Failed after {max_retries + 1} attempts: {last_error}"],
            "success": False,
            "retry_count": max_retries + 1,
        }

    def cleanup(self):
        """Clean up resources."""
        try:
            if hasattr(self, 'text_vdb'):
                self.text_vdb.cleanup()
            if hasattr(self, 'image_vdb'):
                self.image_vdb.cleanup()
        except Exception:
            pass  # Ignore cleanup errors