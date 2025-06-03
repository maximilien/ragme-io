from fastapi import FastAPI, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import Optional, Dict, Any
import PyPDF2
import docx
import tempfile
import os
from pathlib import Path

app = FastAPI(title="MCP PDF Processing Server")

class ToolResponse(BaseModel):
    success: bool
    data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

@app.post("/tool/process_pdf", response_model=ToolResponse)
async def process_pdf(file: UploadFile = File(...)):
    try:
        # Validate file type
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="File must be a PDF")

        # Create a temporary file to store the uploaded PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the PDF using PyPDF2
            with open(temp_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Extract text from all pages
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                # Get metadata
                metadata = pdf_reader.metadata

                return ToolResponse(
                    success=True,
                    data={
                        "data": {
                            "filename": file.filename,
                            "text": text,
                            "page_count": len(pdf_reader.pages)
                        },
                        "metadata": metadata
                    }
                )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(
            success=False,
            error=str(e)
        )

@app.post("/tool/process_docx", response_model=ToolResponse)
async def process_docx(file: UploadFile = File(...)):
    try:
        # Validate file type
        if not file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a DOCX")

        # Create a temporary file to store the uploaded DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            # Process the DOCX using python-docx
            doc = docx.Document(temp_file_path)
            
            # Extract text from all paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)

            # Get core properties
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }

            return ToolResponse(
                success=True,
                data={
                    "data": {
                        "filename": file.filename,
                        "text": text,
                        "tables": tables,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables)
                    },
                    "metadata": metadata
                }
            )
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    except Exception as e:
        return ToolResponse(
            success=False,
            error=str(e)
        )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8010) 